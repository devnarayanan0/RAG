import logging
import io
import base64
from pathlib import Path
from docx import Document
from PIL import Image

logger = logging.getLogger(__name__)

def extract_docx(path: str) -> str:
    """Extract text from DOCX preserving structure."""
    document = Document(path)
    paragraphs = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            paragraphs.append(text)

    return "\n".join(paragraphs)

def extract_paragraph_images(paragraph) -> list:
    """
    Extract embedded images from a python-docx paragraph XML.
    Scans elements for relationship IDs pointing to image parts.
    """
    images = []
    p_element = paragraph._p
    r_ns = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
    
    for elem in p_element.xpath('.//*'):
        for name, value in elem.attrib.items():
            if name.startswith(f"{{{r_ns}}}"):
                rId = value
                try:
                    if rId in paragraph.part.rels:
                        rel = paragraph.part.rels[rId]
                        target_part = getattr(rel, "target_part", None)
                        if target_part:
                            content_type = getattr(target_part, "content_type", "")
                            if "image" in content_type.lower() or "image" in getattr(rel, "target_ref", "").lower():
                                image_data = target_part.blob
                                if not image_data or len(image_data) < 100:
                                    continue
                                
                                img = Image.open(io.BytesIO(image_data))
                                width, height = img.size
                                
                                base64_str = base64.b64encode(image_data).decode('utf-8')
                                
                                if not any(x["size_bytes"] == len(image_data) for x in images):
                                    images.append({
                                        "base64": base64_str,
                                        "width": width,
                                        "height": height,
                                        "mime_type": content_type or "image/png",
                                        "size_bytes": len(image_data)
                                    })
                except Exception as e:
                    logger.debug(f"Error checking docx image relation {rId}: {e}")
    return images

def extract_docx_live(path: str):
    """
    Live DOCX extraction generator.
    Yields logical page data with native text, OCR text (None), and embedded images.
    """
    try:
        doc = Document(path)
        pages_data = []
        current_text = []
        current_images = []
        current_char_count = 0
        
        for paragraph in doc.paragraphs:
            # Check for page break before paragraph
            if paragraph.paragraph_format.page_break_before and (current_text or current_images):
                pages_data.append({
                    "text": "\n".join(current_text),
                    "images": current_images
                })
                current_text = []
                current_images = []
                current_char_count = 0
                
            p_images = extract_paragraph_images(paragraph)
            current_images.extend(p_images)
            
            text = paragraph.text.strip()
            if text:
                current_text.append(text)
                current_char_count += len(text)
                
            # Check for inline page break element
            has_inline_break = False
            for br in paragraph._p.xpath('.//w:br'):
                if br.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type') == 'page':
                    has_inline_break = True
                    break
                    
            if (has_inline_break or current_char_count >= 2000) and (current_text or current_images):
                pages_data.append({
                    "text": "\n".join(current_text),
                    "images": current_images
                })
                current_text = []
                current_images = []
                current_char_count = 0
                
        if current_text or current_images or not pages_data:
            pages_data.append({
                "text": "\n".join(current_text),
                "images": current_images
            })
            
        total_pages = len(pages_data)
        
        for idx, page in enumerate(pages_data, 1):
            yield {
                "page": idx,
                "native_text": page["text"] if page["text"] else None,
                "ocr_text": None,
                "images": page["images"],
                "total_pages": total_pages
            }
            
            total_chars = len(page["text"]) if page["text"] else 0
            logger.info(
                f"Extracted DOCX page {idx}/{total_pages}: "
                f"{total_chars} chars, {len(page['images'])} images"
            )
            
    except Exception as e:
        logger.error(f"DOCX extraction error: {e}")
        raise

def process_docx(path: str) -> dict:
    """Extract and return text with metadata."""
    text = extract_docx(path)
    filename = Path(path).name

    return {
        "text": text,
        "metadata": {
            "source_type": "docx",
            "filename": filename,
            "ocr_used": False,
            "path": path
        }
    }